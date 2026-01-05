from docx import Document
from docx.shared import Inches, Pt
import os
import base64
import tempfile

class DocxManager:
    def __init__(self):
        self.doc = None
        self.doc_path = None

    def create_new(self):
        self.doc = Document()
        self.doc_path = None
        return "New document created."

    def load_document(self, path):
        if not os.path.exists(path):
            raise FileNotFoundError(f"File not found: {path}")
        self.doc = Document(path)
        self.doc_path = path
        return f"Document loaded: {path}"

    def save_document(self, path=None):
        if not self.doc:
            raise ValueError("No document is open.")
        
        target_path = path if path else self.doc_path
        if not target_path:
            raise ValueError("No save path specified.")
            
        self.doc.save(target_path)
        self.doc_path = target_path
        return f"Document saved to {target_path}"

    def get_structure(self):
        if not self.doc:
            raise ValueError("No document is open.")
        
        structure = []
        for element in self.doc.element.body:
            if element.tag.endswith('p'):  # Paragraph
                # Find the paragraph object corresponding to this element
                # This is a bit tricky, doing a simpler iteration for now
                pass
        
        # Simpler approach: iterate paragraphs and tables in order? 
        # python-docx doesn't easily give mixed order. 
        # We will return lists of paragraphs and tables for now, 
        # or try to reconstruct order if needed.
        # For simplicity in V1, let's just dump paragraphs then tables, 
        # or just basic text content.
        
        content = []
        for p in self.doc.paragraphs:
            content.append({
                "type": "paragraph",
                "text": p.text,
                "style": p.style.name
            })
            
        return content
        
    def add_paragraph(self, text, style=None):
        if not self.doc:
            raise ValueError("No document is open.")
        p = self.doc.add_paragraph(text, style=style)
        return f"Paragraph added. Text: '{text[:20]}...'"

    def add_heading(self, text, level=1):
        if not self.doc:
            raise ValueError("No document is open.")
        self.doc.add_heading(text, level=level)
        return f"Heading added. Text: '{text}'"

    def add_table(self, rows, cols, data=None):
        if not self.doc:
            raise ValueError("No document is open.")
        table = self.doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'
        
        if data:
            for i, row_data in enumerate(data):
                if i >= rows: break
                row = table.rows[i]
                for j, cell_text in enumerate(row_data):
                    if j >= cols: break
                    row.cells[j].text = str(cell_text)
                    
        return "Table added."

    def add_image(self, image_path_or_base64, width_inches=None):
        if not self.doc:
            raise ValueError("No document is open.")

        if os.path.exists(image_path_or_base64):
            # It's a file path
            image_path = image_path_or_base64
        else:
            # Assume base64
            try:
                image_data = base64.b64decode(image_path_or_base64)
                with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp:
                    tmp.write(image_data)
                    image_path = tmp.name
            except Exception as e:
                raise ValueError(f"Invalid image input. Provide a valid path or base64 string. Error: {e}")

        width = Inches(width_inches) if width_inches else None
        self.doc.add_picture(image_path, width=width)
        return "Image added."

    def read_full_content(self):
        """
        Returns a simplified representation of the document content including images placeholders.
        """
        if not self.doc:
            raise ValueError("No document is open.")
            
        result = {
            "paragraphs": [p.text for p in self.doc.paragraphs],
            "tables_count": len(self.doc.tables),
            "images_count": self._count_images()
        }
        return result

    def _count_images(self):
        count = 0
        for rel in self.doc.part.rels.values():
            if "image" in rel.target_ref:
                count += 1
        return count

    def extract_images(self, output_dir):
        """Extract all images from the document to the specified directory."""
        if not self.doc:
            raise ValueError("No document is open.")
        
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
            
        image_paths = []
        for rel in self.doc.part.rels.values():
            if "image" in rel.target_ref:
                # This is an image part
                # rel.target_part.blob is the bytes
                # We need a filename. 
                # rel.target_ref often looks like 'media/image1.png'
                
                # Simple extraction
                image_filename = os.path.basename(rel.target_ref)
                image_path = os.path.join(output_dir, image_filename)
                
                with open(image_path, "wb") as f:
                    f.write(rel.target_part.blob)
                
                image_paths.append(image_path)
        
        return image_paths

